"""
Minimal happy-path tests for the SoloPilot requirement analyser.

These tests verify core functionality using simple fixtures.
For comprehensive testing, expand these test cases.
"""

import json
import sys
import tempfile
from pathlib import Path

import pytest

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.agents.analyser.parser import ImageParser, SpecBuilder, TextParser


class TestTextParser:
    """Test TextParser functionality."""

    def test_parse_markdown_file(self):
        """Test parsing a simple markdown file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write(
                """# Project Requirements

## Overview
Build a simple todo app with user authentication.

## Features
- User registration and login
- Create, edit, delete todos
- Mark todos as complete
- Export todos to PDF

## Constraints
- Must use React frontend
- Use Node.js backend
- Mobile responsive design
"""
            )
            f.flush()

            parser = TextParser()
            content = parser.parse_file(f.name)

            assert "Project Requirements" in content
            assert "todo app" in content
            assert "React frontend" in content

            # Clean up
            Path(f.name).unlink()

    def test_parse_text_file(self):
        """Test parsing a plain text file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Simple project brief: Create a weather app that shows current conditions.")
            f.flush()

            parser = TextParser()
            content = parser.parse_file(f.name)

            assert "weather app" in content
            assert "current conditions" in content

            # Clean up
            Path(f.name).unlink()

    def test_parse_docx_file(self):
        """Test parsing a DOCX file."""
        # Create a simple DOCX file for testing
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            # Create a simple DOCX document
            doc = Document()
            doc.add_heading("Project Requirements", 0)
            doc.add_paragraph("Build a task management application.")

            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Feature"
            table.cell(0, 1).text = "Description"
            table.cell(1, 0).text = "User Login"
            table.cell(1, 1).text = "Authentication system"

            doc.add_paragraph("The application must be mobile responsive.")
            doc.save(f.name)

            # Test parsing
            parser = TextParser()
            content = parser.parse_file(f.name)

            assert "Project Requirements" in content
            assert "task management application" in content
            assert "Feature | Description" in content
            assert "User Login | Authentication system" in content
            assert "mobile responsive" in content

            # Clean up
            Path(f.name).unlink()

    def test_unsupported_file_format(self):
        """Test handling of unsupported file formats."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False) as f:
            f.write("test content")
            f.flush()

            parser = TextParser()

            with pytest.raises(ValueError, match="Unsupported file format"):
                parser.parse_file(f.name)

            # Clean up
            Path(f.name).unlink()


class TestImageParser:
    """Test ImageParser functionality."""

    def test_unsupported_image_format(self):
        """Test handling of unsupported image formats."""
        parser = ImageParser()

        with pytest.raises(ValueError, match="Unsupported image format"):
            parser.parse_image("test.xyz")

    def test_batch_parse_empty_list(self):
        """Test batch parsing with empty list."""
        parser = ImageParser()
        results = parser.batch_parse_images([])

        assert isinstance(results, dict)
        assert len(results) == 0


class TestSpecBuilder:
    """Test SpecBuilder functionality."""

    def test_build_basic_specification(self):
        """Test building a basic specification."""
        builder = SpecBuilder()

        text_requirements = {
            "title": "Todo App",
            "summary": "A simple todo application",
            "features": [
                {"name": "User Auth", "desc": "Login and registration"},
                {"name": "Todo CRUD", "desc": "Create, read, update, delete todos"},
            ],
            "constraints": ["Use React", "Mobile responsive"],
            "tech_stack": ["React", "Node.js"],
        }

        image_texts = {}
        file_assets = ["requirements.md"]

        spec = builder.build_specification(text_requirements, image_texts, file_assets)

        # Verify structure
        assert spec["title"] == "Todo App"
        assert spec["summary"] == "A simple todo application"
        assert len(spec["features"]) == 2
        assert "User Auth" in [f["name"] for f in spec["features"]]
        assert spec["constraints"] == ["Use React", "Mobile responsive"]
        assert spec["assets"]["docs"] == ["requirements.md"]
        assert "metadata" in spec
        assert "created_at" in spec["metadata"]

    def test_generate_component_diagram(self):
        """Test component diagram generation."""
        builder = SpecBuilder()

        spec = {
            "features": [
                {"name": "User Management", "desc": "Handle users"},
                {"name": "Todo System", "desc": "Manage todos"},
            ]
        }

        artifacts = builder.generate_artifacts(spec)

        assert "component_diagram" in artifacts
        diagram = artifacts["component_diagram"]
        assert "mermaid" in diagram
        assert "User Management" in diagram
        assert "Todo System" in diagram

    def test_generate_task_flow(self):
        """Test task flow generation."""
        builder = SpecBuilder()

        spec = {
            "features": [
                {"name": "Setup", "desc": "Initial setup"},
                {"name": "Development", "desc": "Core development"},
            ]
        }

        artifacts = builder.generate_artifacts(spec)

        assert "task_flow" in artifacts
        flow = artifacts["task_flow"]
        assert "flowchart TD" in flow
        assert "Setup" in flow
        assert "Development" in flow
        assert "Testing" in flow

    def test_save_artifacts(self):
        """Test saving artifacts to disk."""
        with tempfile.TemporaryDirectory() as temp_dir:
            builder = SpecBuilder(temp_dir)

            spec = {
                "title": "Test Project",
                "summary": "Test summary",
                "features": [],
                "constraints": [],
                "assets": {"images": [], "docs": []},
                "metadata": {"created_at": "2023-01-01T00:00:00", "session_id": "test_session"},
            }

            artifacts = {
                "component_diagram": "```mermaid\ngraph TD\nA --> B\n```",
                "task_flow": "```mermaid\nflowchart TD\nStart --> End\n```",
            }

            session_dir = builder.save_artifacts(spec, artifacts)

            # Verify files were created
            session_path = Path(session_dir)
            assert session_path.exists()
            assert (session_path / "specification.json").exists()
            assert (session_path / "component_diagram.md").exists()
            assert (session_path / "task_flow.md").exists()
            assert (session_path / "README.md").exists()

            # Verify specification content
            with open(session_path / "specification.json") as f:
                saved_spec = json.load(f)
                assert saved_spec["title"] == "Test Project"


class TestIntegration:
    """Integration tests for complete workflows."""

    def test_complete_analysis_workflow(self):
        """Test complete analysis from text to artifacts."""
        with tempfile.TemporaryDirectory() as temp_dir:
            # Create test input file
            input_file = Path(temp_dir) / "requirements.md"
            input_file.write_text(
                """# E-commerce Platform

Build a modern e-commerce platform with the following features:

## Core Features
- Product catalog with search and filtering
- Shopping cart and checkout process
- User accounts and order history
- Admin panel for inventory management

## Technical Requirements
- React frontend with TypeScript
- Node.js backend with Express
- PostgreSQL database
- Payment integration with Stripe
- Mobile responsive design

## Timeline
6-8 weeks development time
"""
            )

            # Create components
            text_parser = TextParser()
            builder = SpecBuilder(temp_dir)

            # Process text
            content = text_parser.parse_file(str(input_file))

            # Note: This test assumes LLM is available
            # In CI/CD, you might want to mock the LLM response
            try:
                requirements = text_parser.extract_requirements(content)

                # Build spec
                spec = builder.build_specification(requirements, {}, [str(input_file)])
                artifacts = builder.generate_artifacts(spec)
                session_dir = builder.save_artifacts(spec, artifacts)

                # Verify output
                session_path = Path(session_dir)
                assert session_path.exists()
                assert (session_path / "specification.json").exists()

                with open(session_path / "specification.json") as f:
                    spec_data = json.load(f)
                    assert (
                        "E-commerce" in spec_data["title"]
                        or "ecommerce" in spec_data["title"].lower()
                    )

            except RuntimeError as e:
                # LLM not available - skip LLM-dependent parts
                if "No LLM available" in str(e):
                    pytest.skip("LLM not available for integration test")
                else:
                    raise


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
