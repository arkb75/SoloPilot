"""
Core parsing classes for the SoloPilot requirement analyser.

This module contains the main components for processing client requirements:
- TextParser: Handles text-based documents (MD, TXT, DOCX)
- ImageParser: Processes images with OCR using pytesseract
- SpecBuilder: Constructs structured JSON specifications and generates artifacts
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pytesseract
import yaml
from PIL import Image

# Load environment variables from .env file
try:
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:
    pass  # dotenv is optional

# Try to import LangChain Bedrock components
try:
    from langchain_aws import ChatBedrock

    LANGCHAIN_AVAILABLE = True
    BEDROCK_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    BEDROCK_AVAILABLE = False
    print("⚠️  LangChain AWS not available. LLM features will be disabled.")

# Import standardized Bedrock client
try:
    from src.common.bedrock_client import (
        BedrockError,
        create_bedrock_client,
        get_standardized_error_message,
    )

    STANDARDIZED_CLIENT_AVAILABLE = True
except ImportError:
    STANDARDIZED_CLIENT_AVAILABLE = False

# Vector similarity - try FAISS first, fallback to scikit-learn
try:
    import faiss  # noqa: F401

    VECTOR_BACKEND = "faiss"
except ImportError:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer  # noqa: F401
        from sklearn.metrics.pairwise import cosine_similarity  # noqa: F401

        VECTOR_BACKEND = "sklearn"
    except ImportError:
        VECTOR_BACKEND = None


class TextParser:
    """Parses text-based client requirements from various formats."""

    def __init__(self, config_path: Optional[str] = None):
        self.config = self._load_config(config_path)
        self.primary_llm = None
        self.fallback_llm = None
        self.standardized_client = None
        self._setup_llm()

    def _load_config(self, config_path: Optional[str]) -> Dict[str, Any]:
        """Load model configuration with environment variable substitution."""
        config_file = config_path or "config/model_config.yaml"

        if os.path.exists(config_file):
            with open(config_file, "r") as f:
                content = f.read()

            # Substitute environment variables in format ${VAR:-default}
            def env_substitute(match):
                var_spec = match.group(1)
                if ":-" in var_spec:
                    var_name, default = var_spec.split(":-", 1)
                    return os.getenv(var_name, default)
                else:
                    return os.getenv(var_spec, "")

            content = re.sub(r"\$\{([^}]+)\}", env_substitute, content)
            return yaml.safe_load(content)

        # Fallback config with inference profile ARN
        return {
            "llm": {
                "primary": "bedrock",
                "bedrock": {
                    "inference_profile_arn": "arn:aws:bedrock:us-east-2:392894085110:inference-profile/us.anthropic.claude-sonnet-4-20250514-v1:0",
                    "region": "us-east-2",
                },
            }
        }

    def _model_id_from_arn(self, arn: str) -> str:
        """Extract modelId from inference profile ARN.

        ARN format: arn:aws:bedrock:<region>:<acct>:inference-profile/<modelId>
        """
        return arn.split("/")[-1]

    def _setup_llm(self):
        """Initialize LLM instances with standardized and legacy support."""
        # Try standardized client first
        if STANDARDIZED_CLIENT_AVAILABLE:
            try:
                self.standardized_client = create_bedrock_client(self.config)
                print("✅ Standardized Bedrock client initialized")
                return
            except BedrockError as e:
                error_msg = get_standardized_error_message(e, "analyser")
                print(error_msg)
                if "NO_NETWORK" in str(e):
                    # Allow offline mode
                    return

        # Fallback to legacy LangChain initialization
        if os.getenv("NO_NETWORK") == "1":
            print("🚫 NO_NETWORK=1, skipping Bedrock initialization")
            return

        if not LANGCHAIN_AVAILABLE:
            return

        # Setup primary LLM (Bedrock) - legacy mode
        if BEDROCK_AVAILABLE and self.config["llm"]["primary"] == "bedrock":
            try:
                bedrock_config = self.config["llm"].get("bedrock", {})
                inference_profile_arn = bedrock_config.get("inference_profile_arn")

                if not inference_profile_arn:
                    print("⚠️  No inference_profile_arn found in config, skipping Bedrock")
                    return

                model_id = self._model_id_from_arn(inference_profile_arn)
                self.primary_llm = ChatBedrock(
                    model_id=model_id,
                    region_name=bedrock_config.get("region", "us-east-2"),
                    model_kwargs=bedrock_config.get(
                        "model_kwargs", {"temperature": 0.1, "max_tokens": 2048}
                    ),
                )
                print("✅ Legacy Bedrock client initialized")
            except Exception as e:
                print(f"⚠️  Legacy Bedrock initialization failed: {e}")

    def parse_file(self, file_path: str) -> str:
        """Parse a single text file and extract content."""
        path = Path(file_path)

        if path.suffix.lower() == ".md":
            return self._parse_markdown(path)
        elif path.suffix.lower() == ".txt":
            return self._parse_text(path)
        elif path.suffix.lower() == ".docx":
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

    def _parse_markdown(self, path: Path) -> str:
        """Parse markdown files."""
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_text(self, path: Path) -> str:
        """Parse plain text files."""
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX files using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx"
            )

        try:
            doc = Document(path)

            # Extract all paragraph text
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX file {path}: {e}")

    def extract_requirements(self, text: str) -> Dict[str, Any]:
        """Extract structured requirements from text using Bedrock LLM."""
        # Try standardized client first
        if self.standardized_client:
            return self._extract_with_standardized_client(text)

        # Fallback to legacy LangChain client
        if LANGCHAIN_AVAILABLE and self.primary_llm:
            return self._extract_with_legacy_client(text)

        # Ultimate fallback to keyword extraction
        print("🔄 No LLM available, using keyword extraction fallback")
        return self._extract_requirements_fallback(text)

    def _extract_with_standardized_client(self, text: str) -> Dict[str, Any]:
        """Extract requirements using standardized Bedrock client."""
        prompt = self._build_extraction_prompt(text)

        try:
            print("🧠 Using standardized Bedrock client")
            response = self.standardized_client.simple_invoke(prompt)
            print("✅ Standardized LLM extraction successful")
            return self._parse_llm_response_text(response)
        except (BedrockError, ValueError, json.JSONDecodeError) as e:
            if isinstance(e, BedrockError):
                error_msg = get_standardized_error_message(e, "analyser")
                print(error_msg)
            else:
                print(f"⚠️  JSON parsing failed: {e}")
            print("🔄 Using keyword extraction fallback")
            return self._extract_requirements_fallback(text)

    def _extract_with_legacy_client(self, text: str) -> Dict[str, Any]:
        """Extract requirements using legacy LangChain client."""
        prompt = self._build_extraction_prompt(text)

        try:
            print(f"🧠 Using legacy LLM: {self.primary_llm.__class__.__name__}")
            response = self.primary_llm.invoke(prompt)
            print("✅ Legacy LLM extraction successful")
            return self._parse_llm_response(response)
        except Exception as e:
            print(f"⚠️  Legacy Bedrock LLM failed ({e})")
            print("🔄 Using keyword extraction fallback")
            return self._extract_requirements_fallback(text)

    def _build_extraction_prompt(self, text: str) -> str:
        """Build prompt for requirement extraction."""
        return f"""
Analyze the following client requirement text and extract structured information.
Return a JSON object with these exact keys:

{{
  "title": "Brief project title",
  "summary": "2-3 sentence project summary",
  "features": [
    {{"name": "Feature name", "desc": "Feature description"}}
  ],
  "constraints": ["List of technical or business constraints"],
  "tech_stack": ["Mentioned technologies or preferences"],
  "timeline": "Estimated timeline if mentioned",
  "budget": "Budget constraints if mentioned"
}}

Client Requirements:
{text}

Respond with ONLY the JSON object, no additional text:
"""

    def _parse_llm_response(self, response) -> Dict[str, Any]:
        """Parse LangChain LLM response into structured data."""
        content = response.content if hasattr(response, "content") else str(response)
        return self._parse_llm_response_text(content)

    def _parse_llm_response_text(self, content: str) -> Dict[str, Any]:
        """Parse raw text response into structured data."""
        # Extract JSON from response
        try:
            # Find JSON block in response
            start = content.find("{")
            end = content.rfind("}") + 1
            if start != -1 and end > start:
                json_str = content[start:end]
                return json.loads(json_str)
        except json.JSONDecodeError:
            pass

        # Fallback: try to parse entire response as JSON
        try:
            return json.loads(content)
        except json.JSONDecodeError:
            raise ValueError("Could not parse LLM response as JSON")

    def _extract_requirements_fallback(self, text: str) -> Dict[str, Any]:
        """Fallback requirement extraction using keyword matching."""
        import re

        text_lower = text.lower()

        # Extract title from headers or first line
        title_match = re.search(r"^#\s+(.+)", text, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
        else:
            # Use first sentence as title
            sentences = text.split(".")
            title = sentences[0][:50] + "..." if len(sentences[0]) > 50 else sentences[0]

        # Extract features from bullet points or numbered lists
        features = []
        feature_patterns = [
            r"[-*]\s+(.+)",  # Bullet points
            r"\d+\.\s+(.+)",  # Numbered lists
            r"##\s+(.+)",  # H2 headers
        ]

        for pattern in feature_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                if len(match.strip()) > 5:  # Filter out very short matches
                    features.append({"name": match.strip()[:30], "desc": match.strip()})

        # Extract constraints from key phrases
        constraints = []
        constraint_keywords = ["must", "should", "required", "constraint", "limitation", "cannot"]
        sentences = text.split(".")

        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in constraint_keywords):
                clean_sentence = sentence.strip()
                if len(clean_sentence) > 10:
                    constraints.append(clean_sentence)

        # Extract tech stack from common technology names
        tech_keywords = [
            "react",
            "angular",
            "vue",
            "node.js",
            "python",
            "java",
            "php",
            "mysql",
            "postgresql",
            "mongodb",
            "redis",
            "docker",
            "kubernetes",
            "aws",
            "azure",
            "gcp",
            "typescript",
            "javascript",
            "html",
            "css",
        ]

        tech_stack = []
        for keyword in tech_keywords:
            if keyword in text_lower:
                tech_stack.append(keyword.title())

        # Extract timeline information
        timeline = None
        timeline_patterns = [r"(\d+)\s+(week|month|day)s?", r"(week|month|day)s?\s*(\d+)"]

        for pattern in timeline_patterns:
            match = re.search(pattern, text_lower)
            if match:
                timeline = match.group(0)
                break

        return {
            "title": title,
            "summary": text[:200] + "..." if len(text) > 200 else text,
            "features": features[:10],  # Limit to 10 features
            "constraints": constraints[:5],  # Limit to 5 constraints
            "tech_stack": list(set(tech_stack)),  # Remove duplicates
            "timeline": timeline,
            "budget": None,  # Difficult to extract without LLM
        }


class ImageParser:
    """Processes images using OCR to extract text requirements."""

    def __init__(self):
        self.supported_formats = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"}

    def parse_image(self, image_path: str) -> str:
        """Extract text from image using OCR."""
        path = Path(image_path)

        if path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported image format: {path.suffix}")

        try:
            image = Image.open(path)
            # Configure pytesseract for better accuracy
            config = "--oem 3 --psm 6"
            text = pytesseract.image_to_string(image, config=config)
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"OCR processing failed: {e}")

    def batch_parse_images(self, image_paths: List[str]) -> Dict[str, str]:
        """Parse multiple images and return text content."""
        results = {}
        for path in image_paths:
            try:
                results[path] = self.parse_image(path)
            except Exception as e:
                results[path] = f"Error: {e}"
        return results


class SpecBuilder:
    """Builds structured specifications and generates helpful artifacts."""

    def __init__(self, output_dir: str = "analysis/output"):
        self.output_dir = Path(output_dir)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.session_dir = self.output_dir / self.timestamp

    def build_specification(
        self, text_requirements: Dict[str, Any], image_texts: Dict[str, str], file_assets: List[str]
    ) -> Dict[str, Any]:
        """Build complete specification from parsed inputs."""

        # Merge image text into requirements if relevant
        combined_text = self._combine_image_text(image_texts)
        if combined_text:
            text_requirements.setdefault("image_content", combined_text)

        # Build final spec
        spec = {
            "title": text_requirements.get("title", "Untitled Project"),
            "summary": text_requirements.get("summary", ""),
            "features": text_requirements.get("features", []),
            "constraints": text_requirements.get("constraints", []),
            "assets": {"images": list(image_texts.keys()), "docs": file_assets},
            "metadata": {
                "created_at": datetime.now().isoformat(),
                "session_id": self.timestamp,
                "tech_stack": text_requirements.get("tech_stack", []),
                "timeline": text_requirements.get("timeline"),
                "budget": text_requirements.get("budget"),
            },
        }

        return spec

    def _combine_image_text(self, image_texts: Dict[str, str]) -> str:
        """Combine text extracted from images."""
        valid_texts = [
            text for text in image_texts.values() if text and not text.startswith("Error:")
        ]
        return "\n\n".join(valid_texts) if valid_texts else ""

    def generate_artifacts(self, spec: Dict[str, Any]) -> Dict[str, str]:
        """Generate helpful artifacts for downstream agents."""
        artifacts = {}

        # Generate component diagram
        artifacts["component_diagram"] = self._generate_component_diagram(spec)

        # Generate task flow
        artifacts["task_flow"] = self._generate_task_flow(spec)

        # Generate optional wireframe
        if self._should_generate_wireframe(spec):
            artifacts["wireframe"] = self._generate_wireframe(spec)

        return artifacts

    def _generate_component_diagram(self, spec: Dict[str, Any]) -> str:
        """Generate Mermaid component diagram."""
        features = spec.get("features", [])

        diagram = "```mermaid\ngraph TD\n"
        diagram += "    A[User Interface] --> B[Business Logic]\n"
        diagram += "    B --> C[Data Layer]\n"

        # Add feature-specific components
        for i, feature in enumerate(features, 1):
            component_id = f"F{i}"
            component_name = feature.get("name", f"Feature {i}")
            diagram += f"    B --> {component_id}[{component_name}]\n"

        diagram += "```"
        return diagram

    def _generate_task_flow(self, spec: Dict[str, Any]) -> str:
        """Generate Mermaid task flow diagram."""
        features = spec.get("features", [])

        flow = "```mermaid\nflowchart TD\n"
        flow += "    Start([Project Start]) --> Setup[Environment Setup]\n"

        for i, feature in enumerate(features, 1):
            feature_id = f"F{i}"
            feature_name = feature.get("name", f"Feature {i}")
            prev_id = f"F{i-1}" if i > 1 else "Setup"
            flow += f"    {prev_id} --> {feature_id}[{feature_name}]\n"

        last_feature = f"F{len(features)}" if features else "Setup"
        flow += f"    {last_feature} --> Test[Testing & QA]\n"
        flow += "    Test --> Deploy[Deployment]\n"
        flow += "    Deploy --> End([Project Complete])\n"
        flow += "```"

        return flow

    def _should_generate_wireframe(self, spec: Dict[str, Any]) -> bool:
        """Determine if wireframe generation is beneficial."""
        # Generate wireframe for UI-heavy projects
        ui_keywords = ["interface", "ui", "frontend", "web", "app", "dashboard"]
        text_content = (
            spec.get("summary", "")
            + " "
            + " ".join(f.get("desc", "") for f in spec.get("features", []))
        ).lower()

        return any(keyword in text_content for keyword in ui_keywords)

    def _generate_wireframe(self, spec: Dict[str, Any]) -> str:
        """Generate ASCII wireframe for UI projects."""
        wireframe = """
ASCII Wireframe (Low-Fidelity):

┌─────────────────────────────────────┐
│              HEADER                 │
│  [Logo]  [Nav1] [Nav2] [Nav3]      │
├─────────────────────────────────────┤
│                                     │
│  Main Content Area                  │
│                                     │
│  ┌─────────────┐ ┌─────────────┐   │
│  │   Feature   │ │   Feature   │   │
│  │    Block    │ │    Block    │   │
│  └─────────────┘ └─────────────┘   │
│                                     │
├─────────────────────────────────────┤
│              FOOTER                 │
└─────────────────────────────────────┘
"""
        return wireframe

    def save_artifacts(self, spec: Dict[str, Any], artifacts: Dict[str, str]) -> str:
        """Save specification and artifacts to disk."""
        # Create session directory
        self.session_dir.mkdir(parents=True, exist_ok=True)

        # Save main specification
        spec_path = self.session_dir / "specification.json"
        with open(spec_path, "w") as f:
            json.dump(spec, f, indent=2)

        # Save artifacts
        for name, content in artifacts.items():
            artifact_path = self.session_dir / f"{name}.md"
            with open(artifact_path, "w") as f:
                f.write(f"# {name.replace('_', ' ').title()}\n\n{content}")

        # Create summary file
        summary_path = self.session_dir / "README.md"
        with open(summary_path, "w") as f:
            f.write(self._generate_session_summary(spec, artifacts))

        return str(self.session_dir)

    def _generate_session_summary(self, spec: Dict[str, Any], artifacts: Dict[str, str]) -> str:
        """Generate session summary README."""
        summary = f"""# Analysis Session: {spec['title']}

**Generated:** {spec['metadata']['created_at']}
**Session ID:** {spec['metadata']['session_id']}

## Summary
{spec['summary']}

## Features
"""
        for feature in spec.get("features", []):
            summary += (
                f"- **{feature.get('name', 'Unnamed')}**: {feature.get('desc', 'No description')}\n"
            )

        if spec.get("constraints"):
            summary += "\n## Constraints\n"
            for constraint in spec["constraints"]:
                summary += f"- {constraint}\n"

        summary += "\n## Generated Artifacts\n"
        for artifact_name in artifacts.keys():
            summary += f"- [{artifact_name.replace('_', ' ').title()}]({artifact_name}.md)\n"

        return summary
